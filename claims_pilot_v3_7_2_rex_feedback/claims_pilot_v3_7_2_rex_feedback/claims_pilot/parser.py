from __future__ import annotations

import io
import re
from dataclasses import dataclass, asdict
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any, Dict, List

from .visual import ImageFinding, analyze_image_bytes, summarize_image_findings


@dataclass
class ParsedFile:
    name: str
    text: str
    image_findings: List[ImageFinding]
    attachments: List[Dict[str, Any]]

    def to_summary(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "chars": len(self.text or ""),
            "images": [f.to_dict() for f in self.image_findings],
            "attachments": self.attachments,
        }


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_eml_bytes_detailed(name: str, data: bytes) -> ParsedFile:
    msg = BytesParser(policy=policy.default).parsebytes(data)
    chunks: List[str] = []
    attachments: List[Dict[str, Any]] = []
    image_findings: List[ImageFinding] = []

    subject = msg.get("subject", "")
    if subject:
        chunks.append(f"Objet: {subject}")
    sender = msg.get("from", "")
    if sender:
        chunks.append(f"De: {sender}")
    date = msg.get("date", "")
    if date:
        chunks.append(f"Date email: {date}")

    # First pass: text content, because image analysis benefits from context.
    if msg.is_multipart():
        for part in msg.walk():
            if part.is_multipart():
                continue
            ctype = part.get_content_type()
            dispo = part.get_content_disposition()
            fn = part.get_filename() or ""
            if ctype in ("text/plain", "text/html") and dispo != "attachment":
                try:
                    payload = part.get_content()
                    if ctype == "text/html":
                        payload = strip_html(str(payload))
                    chunks.append(str(payload))
                except Exception:
                    pass
            elif ctype == "application/pdf" or ctype.startswith("image/") or dispo in {"attachment", "inline"}:
                raw = part.get_payload(decode=True) or b""
                attachments.append({"filename": fn or ctype, "content_type": ctype, "size": len(raw), "disposition": dispo or ""})
    else:
        try:
            payload = msg.get_content()
            if msg.get_content_type() == "text/html":
                payload = strip_html(str(payload))
            chunks.append(str(payload))
        except Exception:
            pass

    base_text = normalize_text("\n\n".join(chunks))

    # Second pass: attachments and inline images.
    if msg.is_multipart():
        for idx, part in enumerate(msg.walk()):
            if part.is_multipart():
                continue
            ctype = part.get_content_type()
            fn = part.get_filename() or f"part_{idx}"
            raw = part.get_payload(decode=True) or b""
            if not raw:
                continue
            if ctype.startswith("image/"):
                image_findings.append(analyze_image_bytes(raw, name, fn, base_text))
            elif ctype == "application/pdf":
                pdf_parsed = parse_pdf_bytes_detailed(fn, raw, context_text=base_text)
                if pdf_parsed.text:
                    chunks.append(f"\n\n=== PIECE JOINTE PDF : {fn} ===\n{pdf_parsed.text}")
                image_findings.extend(pdf_parsed.image_findings)

    final_text = normalize_text("\n\n".join(chunks))
    image_text = summarize_image_findings(image_findings)
    if image_text:
        final_text = normalize_text(final_text + "\n\n" + image_text)
    return ParsedFile(name=name, text=final_text, image_findings=image_findings, attachments=attachments)


def parse_eml_bytes(data: bytes) -> str:
    return parse_eml_bytes_detailed("email.eml", data).text


def strip_html(html: str) -> str:
    html = re.sub(r"<br\s*/?>", "\n", html, flags=re.I)
    html = re.sub(r"</p>", "\n", html, flags=re.I)
    html = re.sub(r"<[^>]+>", " ", html)
    html = html.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return normalize_text(html)



def _ocr_pdf_page_with_tesseract(data: bytes, page_index: int) -> str:
    """OCR optionnel d'une page PDF scannée.

    Utilisé seulement quand l'extraction texte PDF classique est vide ou quasi vide.
    Sur Streamlit Cloud, le fichier packages.txt installe tesseract-ocr et tesseract-ocr-fra.
    En local, si Tesseract n'est pas disponible, la fonction échoue silencieusement.
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except Exception:
        return ""
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        page = doc.load_page(page_index)
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        try:
            txt = pytesseract.image_to_string(img, lang="fra+eng")
        except Exception:
            txt = pytesseract.image_to_string(img)
        return normalize_text(txt)
    except Exception:
        return ""

def parse_pdf_bytes_detailed(name: str, data: bytes, context_text: str = "") -> ParsedFile:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        return ParsedFile(name, f"[Lecture PDF impossible : pypdf non installé - {exc}]", [], [])
    pages: List[str] = []
    image_findings: List[ImageFinding] = []
    try:
        reader = PdfReader(io.BytesIO(data))
        for idx, page in enumerate(reader.pages):
            txt = page.extract_text() or ""
            # OCR optionnel pour les scans : indispensable pour les déclarations faxées/scannées.
            if len(normalize_text(txt)) < 40:
                ocr_txt = _ocr_pdf_page_with_tesseract(data, idx)
                if ocr_txt:
                    txt = (txt + "\n" if txt else "") + "[OCR PAGE SCANNEE]\n" + ocr_txt
            if txt.strip():
                pages.append(f"--- Page {idx + 1} ---\n{txt}")
            try:
                for j, img in enumerate(page.images):
                    raw = img.data
                    img_name = getattr(img, "name", f"page_{idx+1}_image_{j+1}") or f"page_{idx+1}_image_{j+1}"
                    image_findings.append(analyze_image_bytes(raw, name, f"page {idx+1} - {img_name}", context_text + "\n" + txt))
            except Exception:
                pass
        text = normalize_text("\n\n".join(pages))
        if not text:
            text = "[PDF scanné : aucun texte exploitable extrait automatiquement. Activer l'OCR ou fournir une version texte pour fiabiliser la qualification.]"
        image_text = summarize_image_findings(image_findings)
        if image_text:
            text = normalize_text(text + "\n\n" + image_text)
        return ParsedFile(name, text, image_findings, [])
    except Exception as exc:
        return ParsedFile(name, f"[Lecture PDF impossible : {exc}]", image_findings, [])


def parse_pdf_bytes(data: bytes) -> str:
    return parse_pdf_bytes_detailed("document.pdf", data).text


def parse_docx_bytes(data: bytes) -> str:
    try:
        import docx
    except Exception as exc:
        return f"[Lecture DOCX impossible : python-docx non installé - {exc}]"
    try:
        with io.BytesIO(data) as bio:
            d = docx.Document(bio)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
        return normalize_text("\n".join(parts))
    except Exception as exc:
        return f"[Lecture DOCX impossible : {exc}]"


def parse_uploaded_file_detailed(name: str, data: bytes, context_text: str = "") -> ParsedFile:
    suffix = Path(name).suffix.lower()
    if suffix == ".eml":
        return parse_eml_bytes_detailed(name, data)
    if suffix == ".pdf":
        return parse_pdf_bytes_detailed(name, data)
    if suffix in (".jpg", ".jpeg", ".png", ".webp"):
        finding = analyze_image_bytes(data, name, name, context_text)
        return ParsedFile(name, summarize_image_findings([finding]), [finding], [{"filename": name, "content_type": suffix.replace('.', 'image/'), "size": len(data), "disposition": "upload"}])
    if suffix in (".txt", ".md", ".csv"):
        for enc in ("utf-8", "cp1252", "latin-1"):
            try:
                return ParsedFile(name, normalize_text(data.decode(enc)), [], [])
            except Exception:
                pass
        return ParsedFile(name, normalize_text(data.decode("utf-8", errors="ignore")), [], [])
    if suffix == ".docx":
        return ParsedFile(name, parse_docx_bytes(data), [], [])
    return ParsedFile(name, f"[Type de fichier non lu automatiquement : {suffix}]", [], [])


def parse_uploaded_file(name: str, data: bytes) -> str:
    return parse_uploaded_file_detailed(name, data).text
